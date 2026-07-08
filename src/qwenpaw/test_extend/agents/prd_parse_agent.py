# -*- coding: utf-8 -*-
"""PRD Parse Agent - parses requirement documents with real extraction and LLM analysis.

Supports PDF, DOCX, Markdown, and plain text documents.
Uses platform's LLM for structured analysis when available.
"""

import json
import logging
import shutil
from pathlib import Path

from storage.paths import get_prd_dir

logger = logging.getLogger(__name__)


def _extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("PDF extraction failed: %s", e)
        return ""


def _extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        return ""


def _extract_text_from_markdown(file_path: Path) -> str:
    """Extract text from Markdown (strip HTML tags if mixed)."""
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        text = soup.get_text(separator="\n")
        if text.strip():
            return text
    except Exception:
        pass
    return raw


def _extract_text(file_path: Path) -> str:
    """Extract text from a document based on its extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_from_docx(file_path)
    elif suffix in (".md", ".markdown", ".txt", ".text"):
        return _extract_text_from_markdown(file_path)
    else:
        # Try as text
        try:
            return file_path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""


def _analyze_with_rules(text: str) -> dict:
    """Rule-based structured extraction from PRD text (fallback when LLM unavailable)."""
    import re

    lines = text.split("\n")
    business_flows = []
    validation_rules = []
    exception_flows = []
    risk_checklist = []
    functional_modules = []
    non_functional = []

    current_section = ""
    current_flow = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Detect section headers
        lower = stripped.lower()
        if stripped.startswith("#") or stripped.startswith("##") or stripped.startswith("###"):
            current_section = lower
            # Section-based extraction
            if any(k in lower for k in ["flow", "scenario", "用户流程", "业务流程"]):
                if current_flow:
                    business_flows.append(current_flow)
                current_flow = {"name": stripped.lstrip("#").strip(), "steps": []}
            elif any(k in lower for k in ["module", "feature", "功能模块", "功能点"]):
                pass  # modules extracted from content below
            continue

        # Bullet points / numbered items
        item = re.sub(r"^[-*•\d.)\]]+\s*", "", stripped).strip()
        if not item:
            continue

        # Functional modules from feature/module sections (exclude non-functional)
        if any(k in current_section for k in ["module", "feature", "功能"]) and "non-" not in current_section:
            if len(item) < 100:
                functional_modules.append(item)
        elif "function" in current_section and "non-" not in current_section and "non_function" not in current_section:
            if len(item) < 100:
                functional_modules.append(item)

        # Business flow steps
        if current_flow is not None and any(k in current_section for k in ["flow", "scenario", "流程"]):
            current_flow["steps"].append(item)

        # Non-functional requirements (check before validation to avoid overlap)
        if any(k in lower for k in ["performance", "security", "available", "concurrent", "load time", "response time", "uptime", "scalab", "并发", "性能", "安全", "可用", "负载", "响应时间", "用户"]):
            if any(k in lower for k in ["concurrent", "load", "time", "uptime", "scalab", "响应", "负载", "用户数", "handle"]):
                non_functional.append(item)
                continue

        # Validation rules
        if any(k in lower for k in ["must", "should", "require", "valid", "校验", "验证", "必须", "应当"]):
            validation_rules.append(item)

        # Exception/error flows
        if any(k in lower for k in ["error", "exception", "fail", "invalid", "异常", "错误", "失败"]):
            exception_flows.append({"scenario": item, "handling": ""})

        # Risk items
        if any(k in lower for k in ["risk", "concern", "caution", "风险", "注意", "警告"]):
            risk_checklist.append(item)

    if current_flow and current_flow["steps"]:
        business_flows.append(current_flow)

    # If no functional modules found, extract from bullet items in feature sections
    if not functional_modules:
        for line in lines:
            stripped = line.strip()
            lower = stripped.lower()
            item = re.sub(r"^[-*•\d.)\]]+\s*", "", stripped).strip()
            if any(k in lower for k in ["user can", "user should", "support", "提供", "支持"]) and len(item) < 120:
                functional_modules.append(item)

    # Deduplicate
    functional_modules = list(dict.fromkeys(functional_modules))[:20]
    validation_rules = list(dict.fromkeys(validation_rules))[:15]
    exception_flows = [{"scenario": s, "handling": ""} for s in dict.fromkeys(item["scenario"] for item in exception_flows)][:10]
    non_functional = list(dict.fromkeys(non_functional))[:10]
    risk_checklist = list(dict.fromkeys(risk_checklist))[:10]

    return {
        "business_flows": business_flows,
        "validation_rules": validation_rules,
        "exception_flows": exception_flows,
        "risk_checklist": risk_checklist,
        "functional_modules": functional_modules,
        "non_functional_requirements": non_functional,
    }


async def _analyze_with_llm(text: str, model) -> dict:
    """Use platform LLM to extract structured information from PRD text."""
    if not text.strip():
        return {}

    prompt = """You are a senior test analyst. Analyze the following PRD document and extract structured information.

Return ONLY valid JSON (no markdown, no commentary) with this structure:
{
  "business_flows": [{"name": "...", "steps": ["..."]}],
  "validation_rules": ["..."],
  "exception_flows": [{"scenario": "...", "handling": "..."}],
  "risk_checklist": ["..."],
  "functional_modules": ["..."],
  "non_functional_requirements": ["..."]
}

PRD Content:
""" + text[:12000]

    try:
        from agentscope.message import Msg
        messages = [Msg(role="user", content=prompt, name="test_platform")]

        response = await model(messages)
        if hasattr(response, "__aiter__"):
            accumulated = ""
            async for chunk in response:
                text_content = _extract_chunk_text(chunk)
                if text_content:
                    accumulated = text_content
            raw = accumulated
        else:
            raw = _extract_chunk_text(response)

        if raw:
            # Extract JSON from response
            raw = raw.strip()
            if raw.startswith("```"):
                # Remove code fences
                lines = raw.split("\n")
                raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
            result = json.loads(raw)
            return result
    except Exception as e:
        logger.warning("LLM analysis failed: %s", e)

    return {}


def _extract_chunk_text(chunk) -> str:
    """Extract text from a response chunk."""
    if isinstance(chunk, str):
        return chunk
    if isinstance(chunk, dict):
        return chunk.get("text", chunk.get("content", ""))
    text = getattr(chunk, "text", None)
    if text:
        return text
    content = getattr(chunk, "content", None)
    if isinstance(content, str):
        return content
    return ""


class PrdParseAgent:
    """Agent responsible for parsing PRD documents, OpenAPI specs, and Figma links."""

    def __init__(self, workspace_dir: Path):
        self._workspace_dir = workspace_dir

    async def parse_document(
        self,
        file_path: str,
        iteration_id: str,
        file_content: str = "",
    ) -> dict:
        prd_dir = get_prd_dir(self._workspace_dir, iteration_id)
        prd_dir.mkdir(parents=True, exist_ok=True)

        src = Path(file_path)
        raw_text = ""

        # Priority 1: use provided content
        if file_content:
            raw_text = file_content
            # Save content to workspace
            save_path = prd_dir / src.name
            save_path.write_text(file_content, encoding="utf-8")
        elif src.exists():
            shutil.copy2(src, prd_dir / src.name)
            raw_text = _extract_text(src)

        result = {
            "status": "ok",
            "file": src.name,
            "iteration_id": iteration_id,
            "business_flows": [],
            "validation_rules": [],
            "exception_flows": [],
            "risk_checklist": [],
            "functional_modules": [],
            "non_functional_requirements": [],
            "raw_text": raw_text[:5000] if raw_text else "",
        }

        # Analyze with LLM first, fallback to rule-based extraction
        if raw_text.strip():
            llm_success = False
            try:
                from qwenpaw.providers.provider_manager import ProviderManager
                model = ProviderManager.get_active_chat_model()
                llm_result = await _analyze_with_llm(raw_text, model)
                if llm_result:
                    result.update(llm_result)
                    llm_success = True
            except Exception as e:
                logger.info("LLM not available for PRD analysis: %s", e)

            if not llm_success:
                rule_result = _analyze_with_rules(raw_text)
                result.update(rule_result)

        return result

    async def parse_openapi(self, spec_url: str, iteration_id: str) -> dict:
        result = {
            "status": "ok",
            "source": spec_url,
            "iteration_id": iteration_id,
            "endpoints": [],
            "schemas": [],
            "authentication": "unknown",
        }

        # Try to fetch and parse the OpenAPI spec
        if spec_url.startswith(("http://", "https://")):
            try:
                import httpx
                async with httpx.AsyncClient(timeout=30) as client:
                    resp = await client.get(spec_url)
                    spec = resp.json()
                    result["endpoints"] = [
                        {"method": m.upper(), "path": p}
                        for p, methods in spec.get("paths", {}).items()
                        for m in methods
                        if m.upper() in ("GET", "POST", "PUT", "DELETE", "PATCH")
                    ]
                    result["schemas"] = list(spec.get("components", {}).get("schemas", {}).keys())
                    security = spec.get("security", [])
                    if security:
                        result["authentication"] = str(security)
            except Exception as e:
                logger.warning("OpenAPI fetch failed: %s", e)

        return result

    async def parse_figma(self, figma_url: str, iteration_id: str) -> dict:
        return {
            "status": "ok",
            "source": figma_url,
            "iteration_id": iteration_id,
            "screens": [],
            "components": [],
            "flows": [],
        }

    async def identify_ambiguities(self, parsed_prd: dict) -> dict:
        risks = []
        if not parsed_prd.get("validation_rules"):
            risks.append("No validation rules extracted from PRD")
        if not parsed_prd.get("exception_flows"):
            risks.append("No exception flows identified")
        if not parsed_prd.get("business_flows"):
            risks.append("No business flows identified")
        if not parsed_prd.get("functional_modules"):
            risks.append("No functional modules identified")
        return {"risk_checklist": risks, "severity": "medium" if risks else "low"}
